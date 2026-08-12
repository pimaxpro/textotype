import os
import re
import subprocess
import tempfile
import streamlit as st
from docx import Document

st.set_page_config(
    page_title="Tool Chuẩn Hóa LaTeX sang Word",
    page_icon="📝",
    layout="wide"
)

st.title("📝 Tool Chuẩn Hóa LaTeX (ex_test) sang Word Pro")
st.write("Ứng dụng hỗ trợ chuyển đổi đề thi LaTeX sang Word chuẩn đẹp, xóa bỏ hoàn toàn ký tự lạ.")

# ==========================================
# ⚙️ 1. SIDEBAR CẤU HÌNH CÁC CHỨC NĂNG
# ==========================================
st.sidebar.header("⚙️ Cấu hình xuất file")

math_format = st.sidebar.radio(
    "Định dạng công thức toán:",
    ["Word Equation (OMML)", "MathType (Tự động / Toggle TeX)"],
    help="• Word Equation: Công thức hiển thị sẵn ở dạng toán học chuẩn của Word.\n• MathType: Giữ mã $...$ để chuyển thành MathType khi mở Word."
)

clean_ex_test = st.sidebar.checkbox(
    "Xử lý chuẩn hóa gói ex_test", 
    value=True, 
    help="Tự động đánh số Câu 1, Câu 2... và tách rõ các đáp án A, B, C, D."
)

st.sidebar.markdown("---")
st.sidebar.info(
    "💡 **Mẹo cho MathType:** Nếu chọn chế độ MathType, khi mở file Word bạn chỉ cần nhấn **Alt + \\** để toàn bộ công thức tự hóa thành MathType gốc."
)


# ==========================================
# 🛠️ 2. HÀM XỬ LÝ EX_TEST (ĐÃ SỬA LỖI REGEX CHOICE)
# ==========================================
def process_ex_test_content(content: str) -> str:
    """Tự động đếm số câu và định dạng đáp án A, B, C, D an toàn"""
    
    # 1. Đổi môi trường \begin{ex} ... \end{ex}
    cau_counter = 0
    def replace_ex(match):
        nonlocal cau_counter
        cau_counter += 1
        return f"\n\nCâu {cau_counter}. "

    content = re.sub(r'\\begin\{ex\}', replace_ex, content)
    content = re.sub(r'\\end\{ex\}', r'\n', content)
    
    # 2. Đổi môi trường \begin{bt} ... \end{bt}
    bt_counter = 0
    def replace_bt(match):
        nonlocal bt_counter
        bt_counter += 1
        return f"\n\nBài {bt_counter}. "

    content = re.sub(r'\\begin\{bt\}', replace_bt, content)
    content = re.sub(r'\\end\{bt\}', r'\n', content)

    # 3. Xử lý \choice với thuật toán đếm ngoặc {} an toàn 100%
    def parse_choices(text):
        pos = 0
        while True:
            match = re.search(r'\\choice', text[pos:])
            if not match:
                break
            
            start_idx = pos + match.start()
            curr = pos + match.end()
            args = []
            
            # Bóc tách đúng 4 nhóm ngoặc nhọn {...}
            for _ in range(4):
                while curr < len(text) and text[curr].isspace():
                    curr += 1
                if curr < len(text) and text[curr] == '{':
                    depth = 1
                    arg_start = curr + 1
                    curr += 1
                    while curr < len(text) and depth > 0:
                        if text[curr] == '{':
                            depth += 1
                        elif text[curr] == '}':
                            depth -= 1
                        curr += 1
                    args.append(text[arg_start:curr-1])
                else:
                    break
            
            if len(args) == 4:
                # Làm sạch \True trong đáp án
                clean_args = [re.sub(r'\\True\s*', '', arg).strip() for arg in args]
                replacement = f"\n\nA. {clean_args[0]}\n\nB. {clean_args[1]}\n\nC. {clean_args[2]}\n\nD. {clean_args[3]}\n"
                text = text[:start_idx] + replacement + text[curr:]
                pos = start_idx + len(replacement)
            else:
                pos = start_idx + 7
                
        return text

    content = parse_choices(content)
    content = re.sub(r'\\True\s*', '(Đúng) ', content)

    # Bỏ bớt khai báo gói LaTeX dư thừa để Pandoc không bị rối
    content = re.sub(r'\\documentclass.*?\n', '', content)
    content = re.sub(r'\\usepackage.*?\n', '', content)
    content = re.sub(r'\\begin\{document\}', '', content)
    content = re.sub(r'\\end\{document\}', '', content)

    return content


# ==========================================
# 🔄 3. HÀM CHUYỂN ĐỔI LATEX SANG WORD & ĐỊNH DẠNG IN ĐẬM
# ==========================================
def convert_latex_to_word(latex_text: str, math_option: str, fix_ex_test: bool) -> bytes:
    if fix_ex_test:
        latex_text = process_ex_test_content(latex_text)

    is_mathtype = "MathType" in math_option

    # Bảo vệ công thức toán nếu chọn chế độ MathType
    if is_mathtype:
        latex_text = re.sub(r'\$\$(.*?)\$\$', r' MATHBLOCKSTART \1 MATHBLOCKEND ', latex_text, flags=re.DOTALL)
        latex_text = re.sub(r'\$([^\$]+)\$', r' MATHINLINESTART \1 MATHINLINEEND ', latex_text)

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "input.tex")
        output_path = os.path.join(tmpdir, "output.docx")

        with open(input_path, "w", encoding="utf-8") as f:
            f.write(latex_text)

        # Lệnh Pandoc thực thi chuyển đổi sang DOCX
        cmd = ["pandoc", input_path, "-o", output_path]
        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            raise Exception(result.stderr)

        # Đọc file Word xuất ra để xử lý khôi phục công thức và IN ĐẬM
        doc = Document(output_path)
        for p in doc.paragraphs:
            text = p.text
            
            # Khôi phục công thức $...$ cho MathType
            if is_mathtype and ("MATHBLOCKSTART" in text or "MATHINLINESTART" in text):
                text = re.sub(r'MATHBLOCKSTART\s*(.*?)\s*MATHBLOCKEND', r'$$\1$$', text)
                text = re.sub(r'MATHINLINESTART\s*(.*?)\s*MATHINLINEEND', r'$\1$', text)
                p.text = text

            # Tự động tìm và IN ĐẬM các nhãn "Câu X.", "Bài X.", "A.", "B.", "C.", "D."
            pattern = r'^(Câu\s+\d+\.|Bài\s+\d+\.|[A-D]\.)'
            match = re.match(pattern, p.text.strip())
            if match:
                prefix = match.group(1)
                rest = p.text.strip()[len(prefix):]
                p.text = "" # Xóa đoạn text cũ
                
                # Tạo đoạn text mới với prefix được in đậm
                run_bold = p.add_run(prefix)
                run_bold.bold = True
                
                p.add_run(rest)

        doc.save(output_path)

        with open(output_path, "rb") as f:
            return f.read()


# ==========================================
# 🖥️ 4. GIAO DIỆN CHÍNH (EDITOR & UPLOAD)
# ==========================================
tab1, tab2 = st.tabs(["✍️ Editor Nhập Mã LaTeX", "📁 Tải File LaTeX (.tex)"])

latex_input = ""

default_code = r"""\documentclass{article}
\usepackage{ex_test}
\begin{document}

\begin{ex}
    Nghiệm của phương trình $x^2 - 4x + 3 = 0$ là:
    \choice
    {$x = 1; x = 3$}
    {$x = -1; x = -3$}
    {$x = 1; x = -3$}
    {$x = -1; x = 3$}
\end{ex}

\begin{ex}
    Cho hàm số $y = f(x)$ có bảng biến thiên như hình vẽ.
    \choice
    {$a = 2$}
    {$a = 4$}
    {$a = 1$}
    {$a = 0$}
\end{ex}

\end{document}"""

with tab1:
    st.subheader("Nhập hoặc dán mã LaTeX vào đây:")
    latex_input = st.text_area("Khung soạn thảo LaTeX:", value=default_code, height=320)

with tab2:
    st.subheader("Tải file `.tex` từ máy tính:")
    uploaded_file = st.file_uploader("Chọn file LaTeX của bạn", type=["tex"])
    if uploaded_file is not None:
        latex_input = uploaded_file.read().decode("utf-8")
        st.success(f"✅ Đã tải file lên thành công: **{uploaded_file.name}**")


# ==========================================
# 🚀 5. NÚT XỬ LÝ VÀ TẢI FILE VỀ
# ==========================================
st.divider()

if st.button("🚀 Bắt đầu chuyển đổi sang Word", type="primary", use_container_width=True):
    if not latex_input.strip():
        st.warning("⚠️ Vui lòng nhập mã LaTeX vào Editor hoặc tải file .tex lên trước!")
    else:
        with st.spinner("Đang xử lý chuẩn hóa nội dung và định dạng..."):
            try:
                docx_bytes = convert_latex_to_word(latex_input, math_format, clean_ex_test)
                
                st.success("🎉 Chuyển đổi thành công!")
                st.download_button(
                    label="📥 Tải file Word (.docx) về máy",
                    data=docx_bytes,
                    file_name="DeThi_ExTest_Clean.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error("❌ Đã xảy ra lỗi trong quá trình chuyển đổi:")
                st.code(str(e))
